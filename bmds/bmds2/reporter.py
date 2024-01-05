import os
import re
from io import BytesIO

from docx.shared import Inches

from .. import constants, datasets
from ..reporting.footnotes import TableFootnote
from ..reporting.styling import Report
from ..utils import ff as default_float_formatter


class Reporter:
    PORTRAIT_WIDTH = 6.5  # available width to use for tables

    def __init__(self, template=None, styles=None):
        """
        Create a new Microsoft Word document (.docx) reporter object.

        Parameters
        ----------
        template : str, docx.Document instance, or None.
            If str, the path the word template to be used.
            If a docx.Document object, this object is used.
            If None, the default template is used.

        styles : Instance of ReporterStyleGuide, or None.
            Determines which styles to apply for each component in document.
            If None, then the default ReporterStyleGuide is used. Note that if
            a custom template is specified, a ReporterStyleGuide should
            generally also be specified.

        Returns
        -------
        None.
        """

        if template is not None and styles is not None:
            pass
        elif template is None and styles is None:
            report = Report.build_default()
            template = report.document
            styles = report.styles
        else:
            raise ValueError("template and styles most both be specified, or neither")

        self.styles = styles
        self.doc = template

        # remove first paragraph if it's blank
        if len(self.doc.paragraphs) > 0 and self.doc.paragraphs[0].text == "":
            self.doc._body._body.remove(self.doc.paragraphs[0]._p)

    def add_session(
        self,
        session,
        input_dataset=True,
        summary_table=True,
        recommendation_details=True,
        recommended_model=True,
        all_models=False,
    ):
        """
        Add an existing session to a Word report.

        Parameters
        ----------
        session : bmds.Session
            BMDS session to be included in reporting
        input_dataset : bool
            Include input dataset data table
        summary_table : bool
            Include model summary table
        recommendation_details : bool
            Include model recommendation details table
        recommended_model : bool
            Include the recommended model output and dose-response plot, if
            one exists
        all_models : bool
            Include all models output and dose-response plots

        Returns
        -------
        None.
        """

        self.doc.add_paragraph(session.dataset._get_dataset_name(), self.styles.header_1)
        self.doc.add_paragraph(f"BMDS version: {session.version_pretty}")

        if input_dataset:
            self._add_dataset(session)
            self.doc.add_paragraph()

        if summary_table:
            self._add_session_summary_table(session)
            self.doc.add_paragraph()

        if recommendation_details:
            self._add_recommendation_details_table(session)
            self.doc.add_paragraph()

        if recommended_model and all_models:
            self._add_recommended_model(session)
            self._add_all_models(session, except_recommended=True)
            self.doc.add_paragraph()
        elif recommended_model:
            self._add_recommended_model(session)
            self.doc.add_paragraph()
        elif all_models:
            self._add_all_models(session, except_recommended=False)
            self.doc.add_paragraph()

        self.doc.add_page_break()

    def save(self, filename):
        """
        Save document to a file.

        Parameters
        ----------
        filename : str
            The output string filename

        """
        self.doc.save(os.path.expanduser(filename))

    def _set_col_width(self, column, size_in_inches):
        for cell in column.cells:
            cell.width = Inches(size_in_inches)

    def _add_dataset(self, session):
        footnotes = TableFootnote()
        dataset = session.original_dataset

        # only add doses dropped footnote if modeling led to successful results
        dose_dropped_footnote_text = "Dose group removed in BMD modeling session"
        if session.doses_dropped > 0 and session.recommended_model is not None:
            doses_dropped = list(
                reversed(
                    [
                        idx < session.doses_dropped
                        for idx, _ in enumerate(range(dataset.num_dose_groups))
                    ]
                )
            )
        else:
            doses_dropped = [False] * dataset.num_dose_groups

        self.doc.add_paragraph("Input dataset", self.styles.header_2)
        hdr = self.styles.tbl_header
        ff = default_float_formatter

        dose_units_text = dataset._get_dose_units_text()
        response_units_text = dataset._get_response_units_text()

        if isinstance(dataset, datasets.DichotomousDataset):
            tbl = self.doc.add_table(2, dataset.num_dose_groups + 1, style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, style=hdr)
            self._write_cell(
                tbl.cell(1, 0), "Affected / Total (%)" + response_units_text, style=hdr
            )
            for i, vals in enumerate(
                zip(dataset.doses, dataset.incidences, dataset.ns, doses_dropped, strict=True)
            ):
                self._write_cell(tbl.cell(0, i + 1), vals[0])
                if vals[3]:
                    p = tbl.cell(0, i + 1).paragraphs[0]
                    footnotes.add_footnote(p, dose_dropped_footnote_text)
                self._write_cell(
                    tbl.cell(1, i + 1),
                    f"{vals[1]}/{vals[2]}\n({vals[1] / float(vals[2]):.1%})",
                )

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (self.PORTRAIT_WIDTH - 0.75) / dataset.num_dose_groups
                self._set_col_width(col, w)

        elif isinstance(dataset, datasets.ContinuousIndividualDataset):
            tbl = self.doc.add_table(dataset.num_dose_groups + 1, 2, style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, style=hdr)
            self._write_cell(tbl.cell(0, 1), "Responses" + response_units_text, style=hdr)

            for i, vals in enumerate(
                zip(dataset.doses, dataset.get_responses_by_dose(), doses_dropped, strict=True)
            ):
                resps = ", ".join([ff(v) for v in vals[1]])
                self._write_cell(tbl.cell(i + 1, 0), vals[0])
                if vals[2]:
                    p = tbl.cell(i + 1, 0).paragraphs[0]
                    footnotes.add_footnote(p, dose_dropped_footnote_text)
                self._write_cell(tbl.cell(i + 1, 1), resps)

            self._set_col_width(tbl.columns[0], 1.0)
            self._set_col_width(tbl.columns[1], 5.5)

        elif isinstance(dataset, datasets.ContinuousDataset):
            tbl = self.doc.add_table(3, dataset.num_dose_groups + 1, style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), "Dose" + dose_units_text, style=hdr)
            self._write_cell(tbl.cell(1, 0), "N", style=hdr)
            self._write_cell(tbl.cell(2, 0), "Mean ± SD" + response_units_text, style=hdr)
            for i, vals in enumerate(
                zip(
                    dataset.doses,
                    dataset.ns,
                    dataset.means,
                    dataset.stdevs,
                    doses_dropped,
                    strict=True,
                )
            ):
                self._write_cell(tbl.cell(0, i + 1), vals[0])
                if vals[4]:
                    p = tbl.cell(0, i + 1).paragraphs[0]
                    footnotes.add_footnote(p, dose_dropped_footnote_text)
                self._write_cell(tbl.cell(1, i + 1), vals[1])
                self._write_cell(tbl.cell(2, i + 1), f"{ff(vals[2])} ± {ff(vals[3])}")

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (self.PORTRAIT_WIDTH - 0.75) / dataset.num_dose_groups
                self._set_col_width(col, w)

        # write footnote
        if len(footnotes) > 0:
            footnotes.add_footnote_text(self.doc, self.styles.tbl_footnote)

    def _write_cell(self, cell, value, style=None, float_formatter=None):
        if style is None:
            style = self.styles.tbl_body

        if isinstance(value, float):
            if float_formatter is None:
                ff = default_float_formatter
                value = ff(value)

        cell.paragraphs[0].text = str(value)
        cell.paragraphs[0].style = style

    _VARIANCE_FOOTNOTE_TEMPLATE = (
        "{} case presented (BMDS Test 2 p-value = {}, BMDS Test 3 p-value = {})."
    )

    def _get_variance_footnote(self, models):
        text = None
        for model in models:
            if model.has_successfully_executed:
                text = self._VARIANCE_FOOTNOTE_TEMPLATE.format(
                    model.get_variance_model_name(),
                    default_float_formatter(model.output["p_value2"]),
                    default_float_formatter(model.output["p_value3"]),
                )
                break
        if text is None:
            text = "Model variance undetermined"
        return text

    def _get_summary_comments(self, base_session, session):
        if session.recommended_model is None:
            txt = "No model was recommended as a best-fitting model."
            if base_session.doses_dropped > 0:
                txt += " Doses were dropped until there were only 3 remaining dose-groups."
            return txt
        else:
            return "{} recommended as best-fitting model on the basis of the lowest {}.".format(
                session.recommended_model.name, session.recommended_model.recommended_variable
            )

    def _get_session_for_table(self, base_session):
        """
        Only present session for modeling when doses were dropped if it's succesful;
        otherwise show the original modeling session.
        """
        if base_session.recommended_model is None and base_session.doses_dropped > 0:
            return base_session.doses_dropped_sessions[0]
        return base_session

    def _add_session_summary_table(self, base_session):
        self.doc.add_paragraph("Summary table", self.styles.header_2)
        hdr = self.styles.tbl_header
        session = self._get_session_for_table(base_session)
        model_groups = session._group_models()
        footnotes = TableFootnote()

        tbl = self.doc.add_table(len(model_groups) + 2, 6, style=self.styles.table)

        # write headers
        dose_units_text = session.dataset._get_dose_units_text()

        self._write_cell(tbl.cell(0, 0), "Model", style=hdr)
        self._write_cell(tbl.cell(0, 1), "Goodness of fit", style=hdr)
        self._write_cell(tbl.cell(1, 1), "", style=hdr)  # set style
        self._write_cell(tbl.cell(1, 2), "AIC", style=hdr)
        self._write_cell(tbl.cell(0, 3), "BMD" + dose_units_text, style=hdr)
        self._write_cell(tbl.cell(0, 4), "BMDL" + dose_units_text, style=hdr)
        self._write_cell(tbl.cell(0, 5), "Comments", style=hdr)

        p = tbl.cell(1, 1).paragraphs[0]
        p.add_run("p").italic = True
        p.add_run("-value")

        # add variance footnote to table header if appropriate
        if session.dtype in constants.CONTINUOUS_DTYPES:
            p = tbl.cell(0, 0).paragraphs[0]
            footnotes.add_footnote(p, self._get_variance_footnote(session.models))

        # merge header columns
        tbl.cell(0, 0).merge(tbl.cell(1, 0))
        tbl.cell(0, 1).merge(tbl.cell(0, 2))
        tbl.cell(0, 3).merge(tbl.cell(1, 3))
        tbl.cell(0, 4).merge(tbl.cell(1, 4))
        tbl.cell(0, 5).merge(tbl.cell(1, 5))

        # write body
        for i, model_group in enumerate(model_groups):
            idx = i + 2
            model = model_group[0]
            output = getattr(model, "output", {})
            self._write_cell(tbl.cell(idx, 0), "")  # temp; set style
            self._write_cell(tbl.cell(idx, 1), output.get("p_value4", "-"))
            self._write_cell(tbl.cell(idx, 2), output.get("AIC", "-"))
            self._write_cell(tbl.cell(idx, 3), output.get("BMD", "-"))
            self._write_cell(tbl.cell(idx, 4), output.get("BMDL", "-"))

            self._write_model_name(tbl.cell(idx, 0).paragraphs[0], model_group, footnotes)

        # write comments
        self._write_cell(tbl.cell(2, 5), self._get_summary_comments(base_session, session))
        tbl.cell(2, 5).merge(tbl.cell(len(model_groups) + 1, 5))

        # set column width
        widths = [1.75, 0.8, 0.8, 0.7, 0.7, 1.75]
        for width, col in zip(widths, tbl.columns, strict=True):
            self._set_col_width(col, width)

        # write footnote
        if len(footnotes) > 0:
            footnotes.add_footnote_text(self.doc, self.styles.tbl_footnote)

    def _add_recommendation_details_table(self, base_session):
        self.doc.add_paragraph("Model recommendation details", self.styles.header_2)
        hdr = self.styles.tbl_header
        session = self._get_session_for_table(base_session)
        model_groups = session._group_models()
        footnotes = TableFootnote()

        tbl = self.doc.add_table(len(model_groups) + 1, 3, style=self.styles.table)

        # write headers
        self._write_cell(tbl.cell(0, 0), "Model", style=hdr)
        self._write_cell(tbl.cell(0, 1), "Bin", style=hdr)
        self._write_cell(tbl.cell(0, 2), "Notes", style=hdr)

        def write_warnings(cell, txt, notes):
            p = cell.add_paragraph("")
            p.add_run(txt).bold = True

            if len(notes) == 0:
                cell.add_paragraph("  -")
            else:
                for note in notes:
                    cell.add_paragraph(f"• {note}")

        # write body
        for i, model_group in enumerate(model_groups):
            idx = i + 1
            model = model_group[0]
            bin = model.get_logic_bin_text().title()
            self._write_cell(tbl.cell(idx, 0), "")  # temp; set style
            self._write_model_name(tbl.cell(idx, 0).paragraphs[0], model_group, footnotes)
            self._write_cell(tbl.cell(idx, 1), bin)

            cell = tbl.cell(idx, 2)
            if not model.logic_notes[0] and not model.logic_notes[1] and not model.logic_notes[2]:
                cell.paragraphs[0].text = "-"
            else:
                cell._element.remove(cell.paragraphs[0]._p)
                if model.logic_notes[2]:
                    write_warnings(cell, "Failures", model.logic_notes[2])
                if model.logic_notes[1]:
                    write_warnings(cell, "Warnings", model.logic_notes[1])
                if model.logic_notes[0]:
                    write_warnings(cell, "Cautions", model.logic_notes[0])

            for p in cell.paragraphs:
                p.style = self.styles.tbl_body

        # set column width
        widths = [1.75, 0.75, 4]
        for width, col in zip(widths, tbl.columns, strict=True):
            self._set_col_width(col, width)

        # write footnote
        if len(footnotes) > 0:
            footnotes.add_footnote_text(self.doc, self.styles.tbl_footnote)

    def _write_model_name(self, p, model_group, footnotes):
        def pretty(name):
            name = name.replace("-", " ")
            return re.sub(r"( \d+)", r"\g<1>°", name)

        def collapse_names(models):
            names = [model.name for model in models]
            search_phrases = ["Polynomial-", "Multistage-Cancer-", "Multistage-"]
            for phrase in search_phrases:
                if phrase in "".join(names):
                    full_phrase = [name for name in names if phrase not in name]
                    matches = [name for name in names if phrase in name]
                    full_phrase.append(matches[0])
                    if len(matches) > 1:
                        remainders = ", ".join(matches[1:]).replace(phrase, "")
                        full_phrase.append(remainders)
                    names = full_phrase

            return ", ".join(names)

        p.add_run(pretty(model_group[0].name))

        if model_group[0].recommended:
            footnotes.add_footnote(p, "Recommended model")

        if len(model_group) > 1:
            names = pretty(collapse_names(model_group[1:]))
            p.add_run(f" (equivalent models include {names})")

    def _model_to_docx(self, model):
        if model.has_successfully_executed:
            # print figure
            fig = model.plot()
            with BytesIO() as f:
                fig.savefig(f)
                self.doc.add_picture(f, width=Inches(6))
            fig.clf()

            # print output file
            self.doc.add_paragraph(model.outfile, style=self.styles.fixed_width)
        else:
            self.doc.add_paragraph("No .OUT file was created.")

    def _add_recommended_model(self, session):
        self.doc.add_paragraph("Recommended model", self.styles.header_2)
        if hasattr(session, "recommended_model") and session.recommended_model is not None:
            self._model_to_docx(session.recommended_model)
        else:
            p = self.doc.add_paragraph()
            p.add_run("No model was recommended as a best-fitting model.").italic = True

    def _add_all_models(self, session, except_recommended=False):
        if except_recommended:
            self.doc.add_paragraph("All other models", self.styles.header_2)
        else:
            self.doc.add_paragraph("All model outputs", self.styles.header_2)

        for model in session.models:
            if model.recommended and except_recommended:
                continue

            self._model_to_docx(model)
