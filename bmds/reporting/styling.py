from io import BytesIO
from pathlib import Path
from typing import Any, Self

from docx import Document
from docx.shared import Inches
from pydantic import BaseModel

from ..constants import BMDS_BLANK_VALUE
from ..plotting import close_figure
from ..utils import ff


class ReporterStyleGuide(BaseModel):
    portrait_width: float = 6.5
    table: str = "bmdsTbl"
    tbl_header: str = "bmdsTblHeader"
    tbl_body: str = "bmdsTblBody"
    tbl_footnote: str = "bmdsTblFootnote"
    fixed_width: str = "bmdsOutputFile"
    header_1: str = "Heading 1"
    header_2: str = "Heading 2"
    header_3: str = "Heading 3"
    header_4: str = "Heading 4"

    def get_header_style(self, level: int) -> str:
        return getattr(self, f"header_{level}")


class Report(BaseModel):
    document: Any = None
    styles: ReporterStyleGuide

    @classmethod
    def build_default(cls) -> Self:
        fn = Path(__file__).parent / "templates/base.docx"
        doc = Document(str(fn))
        return Report(document=doc, styles=ReporterStyleGuide())


def write_cell(cell, value, style, formatter=ff):
    if value == BMDS_BLANK_VALUE:
        value = "-"
    elif isinstance(value, float):
        value = formatter(value)
    cell.paragraphs[0].text = str(value)
    cell.paragraphs[0].style = style


def set_column_width(column, size_in_inches: float):
    for cell in column.cells:
        cell.width = Inches(size_in_inches)


def add_mpl_figure(document, fig, size_in_inches: float):
    with BytesIO() as f:
        fig.savefig(f)
        document.add_picture(f, width=Inches(size_in_inches))
    fig.clf()
    close_figure(fig)
