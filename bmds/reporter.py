from collections import namedtuple, OrderedDict
from io import BytesIO
import docx
from docx.shared import Inches
import numpy as np
import os
import re

from . import constants, datasets


StyleGuide = namedtuple('StyleGuide', [
    'table',
    'tbl_header',
    'tbl_body',
    'tbl_footnote',
    'outfile',
    'header_level'
])


def default_float_formatter(value):
    if np.isclose(value, 0.):
        tmpl = '{}'
    elif abs(value) < 0.001:
        tmpl = '{:.1E}'
    else:
        tmpl = '{:.3f}'
    return tmpl.format(value)


class TableFootnote(OrderedDict):
    def __init__(self):
        super().__init__()
        self.ascii_char = 96

    def add_footnote(self, p, text):
        if text not in self:
            self.ascii_char += 1
            self[text] = chr(self.ascii_char)
        self._add_footnote_character(p, self[text])

    def _add_footnote_character(self, p, symbol):
        run = p.add_run(symbol)
        run.font.superscript = True

    def add_footnote_text(self, p):
        for text, char in self.items():
            self._add_footnote_character(p, char)
            p.add_run(' {}\n'.format(text))


class Reporter:

    def __init__(self, template=None, styles=None):
        """
        Create a new Microsoft Word document (.docx) reporter object.

        Parameters
        ----------
        template : str, docx.Document instance, or None.
            If str, the path the word template to be used.
            If a docx.Document object, this object is used.
            If None, the default template is used.

        styles : Instance of StyleGuide, or None.
            Determines which styles to apply for each component in document.
            If None, then the default StyleGuide is used. Note that if a custom
            template is specified, a StyleGuide should generally also be
            specified.

        Returns
        -------
        None.
        """

        if template is None:
            template = os.path.join(
                os.path.dirname(__file__),
                'templates/base.docx'
            )

        if styles is None:
            styles = StyleGuide('bmdsTbl', 'bmdsTblHeader', 'bmdsTblBody',
                                'bmdsTblFootnote', 'bmdsOutputFile', 1)

        self.styles = styles
        self.doc = docx.Document(template)

    def add_session(self, session, title=None,
                    input_dataset=True, summary_table=True,
                    recommended_model=True, all_models=False):
        """
        Add an existing session to a Word report.

        Parameters
        ----------
        session : bmds.Session
            BMDS session to be included in reporting
        title : str or None
            Title to be used to refer to the session. If one is not provided, a
            default value is used.
        input_dataset : bool
            Include input dataset data table
        summary_table : bool
            Include model summary table
        recommended_model : bool
            Include the recommended model output and dose-response plot, if
            one exists
        all_models : bool
            Include all models output and dose-response plots

        Returns
        -------
        None.
        """

        if title is None:
            title = 'BMDS output results'

        self.doc.add_heading(title, self.styles.header_level)

        if input_dataset:
            self._add_dataset(session.dataset)

        if summary_table:
            self._add_session_summary_table(session)

        if recommended_model and all_models:
            self._add_recommended_model(session)
            self._add_all_models(session, except_recommended=True)
        elif recommended_model:
            self._add_recommended_model(session)
        elif all_models:
            self._add_all_models(session, except_recommended=False)

    def save(self, filename):
        """
        Save document to a file.

        Parameters
        ----------
        filename : str
            The output string filename

        """
        self.doc.save(os.path.expanduser(filename))

    def _set_col_width(self, column, size_in_inches):
        for cell in column.cells:
            cell.width = Inches(size_in_inches)

    def _add_dataset(self, dataset):

        self.doc.add_heading('Input dataset', self.styles.header_level + 1)
        hdr = self.styles.tbl_header

        if isinstance(dataset, datasets.DichotomousDataset):

            tbl = self.doc.add_table(2, dataset.num_dose_groups + 1,
                                     style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), 'Dose', style=hdr)
            self._write_cell(tbl.cell(1, 0), 'Response', style=hdr)
            for i, vals in enumerate(zip(dataset.doses,
                                         dataset.incidences,
                                         dataset.ns)):
                self._write_cell(tbl.cell(0, i + 1), vals[0])
                self._write_cell(tbl.cell(1, i + 1),
                                 '{}/{}'.format(vals[1], vals[2]))

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (6.5 - 0.75) / dataset.num_dose_groups
                self._set_col_width(col, w)

        elif isinstance(dataset, datasets.ContinuousIndividualDataset):

            tbl = self.doc.add_table(dataset.num_dose_groups + 1, 2,
                                     style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), 'Dose', style=hdr)
            self._write_cell(tbl.cell(0, 1), 'Responses', style=hdr)

            for i, vals in enumerate(zip(dataset.doses,
                                         dataset.get_responses_by_dose())):
                resps = ', '.join([str(v) for v in vals[1]])
                self._write_cell(tbl.cell(i + 1, 0), vals[0])
                self._write_cell(tbl.cell(i + 1, 1), resps)

            self._set_col_width(tbl.columns[0], 1.0)
            self._set_col_width(tbl.columns[1], 5.5)

        elif isinstance(dataset, datasets.ContinuousDataset):

            tbl = self.doc.add_table(3, dataset.num_dose_groups + 1,
                                     style=self.styles.table)

            self._write_cell(tbl.cell(0, 0), 'Dose', style=hdr)
            self._write_cell(tbl.cell(1, 0), 'N', style=hdr)
            self._write_cell(tbl.cell(2, 0), 'Mean ± SD', style=hdr)
            for i, vals in enumerate(zip(dataset.doses,
                                         dataset.ns,
                                         dataset.means,
                                         dataset.stdevs)):
                self._write_cell(tbl.cell(0, i + 1), vals[0])
                self._write_cell(tbl.cell(1, i + 1), vals[1])
                self._write_cell(tbl.cell(2, i + 1),
                                 '{} ± {}'.format(vals[2], vals[3]))

            for i, col in enumerate(tbl.columns):
                w = 0.75 if i == 0 else (6.5 - 0.75) / dataset.num_dose_groups
                self._set_col_width(col, w)

    def _write_cell(self, cell, value, style=None, float_formatter=None):

        if style is None:
            style = self.styles.tbl_body

        if isinstance(value, float):
            if float_formatter is None:
                float_formatter = default_float_formatter
                value = float_formatter(value)

        cell.paragraphs[0].text = str(value)
        cell.paragraphs[0].style = style

    _VARIANCE_FOOTNOTE_TEMPLATE = '{} case presented (BMDS Test 2 p-value = {}, BMDS Test 3 p-value = {}).'  # noqa

    def _get_variance_footnote(self, models):
        text = None
        for model in models:
            if model.has_successfully_executed:
                text = self._VARIANCE_FOOTNOTE_TEMPLATE.format(
                    model.get_variance_model_name(),
                    default_float_formatter(model.output['p_value2']),
                    default_float_formatter(model.output['p_value3']),
                )
                break
        if text is None:
            text = 'Model variance undetermined'
        return text

    def _get_summary_comments(self, session):
        if session.recommended_model is None:
            return 'No model was recommended as a best-fitting model.'
        else:
            return '{} recommended as best-fitting model on the basis of the lowest {}.'.format(
                session.recommended_model.name,
                session.recommended_model.recommended_variable
            )

    def _add_session_summary_table(self, session):
        self.doc.add_heading('Summary table', self.styles.header_level + 1)
        hdr = self.styles.tbl_header
        model_groups = session._group_models()
        footnotes = TableFootnote()

        tbl = self.doc.add_table(len(model_groups) + 2, 6,
                                 style=self.styles.table)

        # write headers
        self._write_cell(tbl.cell(0, 0), 'Model', style=hdr)
        self._write_cell(tbl.cell(0, 1), 'Goodness of fit', style=hdr)
        self._write_cell(tbl.cell(1, 1), '', style=hdr)  # set style
        self._write_cell(tbl.cell(1, 2), 'AIC', style=hdr)
        self._write_cell(tbl.cell(0, 3), 'BMD', style=hdr)
        self._write_cell(tbl.cell(0, 4), 'BMDL', style=hdr)
        self._write_cell(tbl.cell(0, 5), 'Comments', style=hdr)

        p = tbl.cell(1, 1).paragraphs[0]
        p.add_run('p').italic = True
        p.add_run('-value')

        # add variance footnote to table header if appropriate
        if session.dtype in constants.CONTINUOUS_DTYPES:
            p = tbl.cell(0, 0).paragraphs[0]
            footnotes.add_footnote(p, self._get_variance_footnote(session.models))

        # merge header columns
        tbl.cell(0, 0).merge(tbl.cell(1, 0))
        tbl.cell(0, 1).merge(tbl.cell(0, 2))
        tbl.cell(0, 3).merge(tbl.cell(1, 3))
        tbl.cell(0, 4).merge(tbl.cell(1, 4))
        tbl.cell(0, 5).merge(tbl.cell(1, 5))

        # write body
        for i, model_group in enumerate(model_groups):
            self._to_docx_summary_row(model_group, tbl, i + 2, footnotes)

        # write comments
        self._write_cell(tbl.cell(2, 5), self._get_summary_comments(session))
        tbl.cell(2, 5).merge(tbl.cell(len(model_groups) + 1, 5))

        # set column width
        widths = [1.75, 0.8, 0.8, 0.7, 0.7, 1.75]
        for width, col in zip(widths, tbl.columns):
            self._set_col_width(col, width)

        # write footnote
        if len(footnotes) > 0:
            p = self.doc.add_paragraph(style=self.styles.tbl_footnote)
            footnotes.add_footnote_text(p)

    def _to_docx_summary_row(self, model_group, tbl, idx, footnotes):

        model = model_group[0]
        output = getattr(model, 'output', {})
        self._write_cell(tbl.cell(idx, 0), '')  # temp; set style
        self._write_cell(tbl.cell(idx, 1), output.get('p_value4', '-'))
        self._write_cell(tbl.cell(idx, 2), output.get('AIC', '-'))
        self._write_cell(tbl.cell(idx, 3), output.get('BMD', '-'))
        self._write_cell(tbl.cell(idx, 4), output.get('BMDL', '-'))

        p = tbl.cell(idx, 0).paragraphs[0]
        p.add_run(self._pretty_names(model_group[0].name))
        if model_group[0].recommended:
            footnotes.add_footnote(p, 'Recommended model')
        if len(model_group) > 1:
            p.add_run(' (equivalent models include {})'.format(
                self._get_collapsed_model_names(model_group[1:])))

    def _pretty_names(self, name):
        name = name.replace('-', ' ')
        return re.sub(r'( \d+)', '\g<1>°', name)

    def _get_collapsed_model_names(self, models):
        names = [model.name for model in models]

        search_phrases = ['Polynomial-', 'Multistage-Cancer-', 'Multistage-']
        for phrase in search_phrases:
            if phrase in ''.join(names):
                not_poly = [name for name in names if phrase not in name]
                poly = [name for name in names if phrase in name]
                rem_poly = ', '.join(poly[1:]).replace(phrase, '')
                not_poly.append(poly[0])
                not_poly.append(rem_poly)
                names = not_poly

        return self._pretty_names(', '.join(names))

    def _model_to_docx(self, model):
        if model.has_successfully_executed:
            # print figure
            fig = model.plot()
            with BytesIO() as f:
                fig.savefig(f)
                self.doc.add_picture(f, width=Inches(6))
            fig.clf()

            # print output file
            self.doc.add_paragraph(model.outfile, style=self.styles.outfile)
        else:
            self.doc.add_paragraph('No .OUT file was created.')

    def _add_recommended_model(self, session):
        self.doc.add_heading('Recommended model', self.styles.header_level + 1)
        if hasattr(session, 'recommended_model') and \
                session.recommended_model is not None:
            self._model_to_docx(session.recommended_model)
        else:
            p = self.doc.add_paragraph()
            p.add_run('No model was recommended as a best-fitting model.')\
             .italic = True

    def _add_all_models(self, session, except_recommended=False):
        if except_recommended:
            self.doc.add_heading('All other models',
                                 self.styles.header_level + 1)
        else:
            self.doc.add_heading('All model outputs',
                                 self.styles.header_level + 1)

        for model in session.models:

            if model.recommended and except_recommended:
                continue

            self._model_to_docx(model)
